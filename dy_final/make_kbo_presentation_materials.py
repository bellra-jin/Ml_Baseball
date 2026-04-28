# -*- coding: utf-8 -*-
"""Create presentation materials for the KBO postseason prediction notebook."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(r"C:\Users\Admin\Documents\GitHub\Ml_Baseball\dy_final")
OUT_DIR = Path(r"C:\Users\Admin\Documents\Codex\2026-04-28\files-mentioned-by-the-user-data")
MD_PATH = OUT_DIR / "KBO_가을야구_예측_발표자료.md"
SCRIPT_PATH = OUT_DIR / "KBO_가을야구_예측_발표대본.md"
DOCX_PATH = OUT_DIR / "KBO_가을야구_예측_설명자료.docx"


def read_csv_by_keyword(keyword: str) -> pd.DataFrame:
    matches = [p for p in ROOT.glob("*.csv") if keyword in p.name]
    if not matches:
        matches = [p for p in (ROOT / "data_").glob("*.csv") if keyword in p.name]
    if not matches:
        raise FileNotFoundError(keyword)
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return pd.read_csv(matches[0], encoding=enc)
        except Exception:
            continue
    raise ValueError(matches[0])


pred = read_csv_by_keyword("예측결과")
team = read_csv_by_keyword("team_master")
league_bat = pd.read_csv(ROOT / "data_" / "리그_타격환경.csv", encoding="utf-8-sig")
league_pit = pd.read_csv(ROOT / "data_" / "리그_투구환경.csv", encoding="utf-8-sig")

unnamed = [c for c in pred.columns if str(c).startswith("Unnamed") or str(c) == ""]
pred = pred.drop(columns=unnamed, errors="ignore")

top4 = pred.head(4).copy()
bubble = pred.iloc[4].copy()
kia = pred[pred["팀명"] == "KIA"].iloc[0]
hanwha = pred[pred["팀명"] == "한화"].iloc[0]
bottom = pred.tail(3).copy()

feature_lines = [
    ("승률", "현재 성적의 기본 체력. 순위표의 가장 직접적인 신호"),
    ("원정승률", "홈 어드밴티지 없이도 버티는 힘. 상위권 지속성 판단"),
    ("홈승률", "홈 경기에서 안정적으로 승수를 쌓는 능력"),
    ("ERA/WHIP", "실점 억제력과 출루 허용. 단기 순위보다 팀 안정성을 설명"),
    ("OPS", "타격 생산성. 단순 타율보다 득점 기대를 잘 반영"),
    ("득실차/기대승률", "현재 승률이 실력인지 운인지 보정하는 지표"),
]


def fmt_pct(v) -> str:
    return f"{float(v):.1f}%"


def fmt_rate(v) -> str:
    return f"{float(v):.3f}".replace("0.", ".")


def prediction_table_md() -> str:
    cols = ["팀명", "순위", "승", "패", "승률", "ERA", "OPS", "WHIP", "가을야구_확률(%)", "예측"]
    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join(["---"] * len(cols)) + " |"
    rows = []
    for _, r in pred[cols].iterrows():
        rows.append(
            "| "
            + " | ".join(
                [
                    str(r["팀명"]),
                    str(int(r["순위"])),
                    str(int(r["승"])),
                    str(int(r["패"])),
                    fmt_rate(r["승률"]),
                    f"{float(r['ERA']):.2f}",
                    fmt_rate(r["OPS"]),
                    f"{float(r['WHIP']):.2f}",
                    fmt_pct(r["가을야구_확률(%)"]),
                    str(r["예측"]),
                ]
            )
            + " |"
        )
    return "\n".join([header, sep, *rows])


md = f"""# KBO 2026 가을야구 진출 예측 발표자료

## 1. 한 줄 소개

KBO 2022~2026 정규시즌 데이터를 전처리해 팀별 전력 지표를 통합하고, 2026년 4월 24일 기준 가을야구 진출 가능성을 앙상블 ML 모델로 예측했다.

## 2. 분석 목적

- 현재 순위만으로는 최종 5강 진출 여부를 설명하기 어렵다.
- 팀 타격, 투수, 수비, 주루, 홈/원정 성적, 득실차를 함께 봐야 한다.
- 목표 변수는 `가을야구`, 즉 최종 순위 5위 이내 여부다.

## 3. 사용 데이터

- 기간: 2022~2026 KBO 정규시즌
- 주요 테이블: 팀 순위, 팀 타자, 팀 투수, 팀 수비, 팀 주루, 선수 타자/투수 마스터, 선수 이동 현황
- 핵심 산출물: `team_master_2022_2026.csv`, `2026_가을야구_예측결과.csv`

## 4. 전처리 핵심

- 팀명 표준화: 구단명 표기 차이를 하나로 통일
- 숫자형 변환: 승률, ERA, OPS, WHIP, 홈/원정 기록 등을 분석 가능한 숫자로 변환
- 투수 이닝 변환: `23 1/3` 같은 이닝 표기를 소수형으로 변환
- 파생 지표 생성: OPS, BABIP, ISO, K/9, BB/9, FIP, 득실차, 기대승률, 승률_운
- 팀 단위 통합: 순위 + 타격 + 투수 + 수비 + 주루 데이터를 `team_master`로 결합

## 5. 모델 구조

- Logistic Regression
- Random Forest
- Gradient Boosting
- 세 모델의 예측 확률을 평균내는 앙상블 방식
- 학습: 2022~2025 완료 시즌
- 예측: 2026 시즌 진행중 데이터

## 6. 모델이 본 주요 신호

| 지표 | 발표 포인트 |
| --- | --- |
{chr(10).join(f"| {name} | {desc} |" for name, desc in feature_lines)}

## 7. 최종 예측 결과

{prediction_table_md()}

## 8. 핵심 인사이트

### 인사이트 1. 2026 예측 진출권은 KT, LG, SSG, 삼성

노트북 결과 기준 예측 진출 팀은 {", ".join(top4["팀명"].tolist())}이다. 네 팀 모두 현재 1~4위권이며, 승률과 투수/타격 지표가 함께 양호하게 나왔다.

### 인사이트 2. KIA는 현재 5위지만 모델상 위험 신호

KIA는 현재 {int(kia["순위"])}위지만 예측 확률은 {fmt_pct(kia["가을야구_확률(%)"])}이다. 현재 순위만 보면 5강권이지만 ERA {float(kia["ERA"]):.2f}, WHIP {float(kia["WHIP"]):.2f} 등 세부 지표가 약해 낮게 평가됐다.

### 인사이트 3. 한화는 미진출 예측이지만 경계선 팀

한화는 현재 {int(hanwha["순위"])}위, 확률 {fmt_pct(hanwha["가을야구_확률(%)"])}로 미진출 팀 중 가장 높다. 완전한 탈락권보다는 반등 가능성이 남은 경계선 팀으로 해석할 수 있다.

### 인사이트 4. 4월 순위만으로는 부족하다

노트북 분석에서 4월 순위 단독 가을야구 예측 정확도는 약 60%였다. 그래서 현재 순위에 의존하지 않고, OPS, ERA, WHIP, 홈/원정 승률, 기대승률 등을 함께 반영했다.

## 9. 발표 결론

이 프로젝트의 핵심은 “현재 순위표를 보완하는 세부 전력 분석”이다. 현재 5위라고 해서 반드시 안정적인 5강 팀은 아니며, 5위 밖이어도 지표상 추격 가능성이 남을 수 있다. 따라서 가을야구 가능성은 순위, 승률, 투수력, 공격 생산성, 홈/원정 균형, 득실차를 함께 봐야 한다.

## 10. 한계

- 2026 데이터는 4월 24일 기준 초반 표본이다.
- 시즌 중 부상, 외국인 선수 교체, 트레이드 등 외부 변수는 모델에 충분히 반영되지 않았다.
- 예측 확률은 확정 결과가 아니라 현재 데이터 기반 가능성이다.
"""


script = f"""# KBO 가을야구 예측 발표 대본

## 오프닝

안녕하세요. 저는 KBO 2022년부터 2026년까지의 정규시즌 데이터를 활용해서, 2026년 가을야구 진출 가능성이 높은 팀을 예측하는 분석을 진행했습니다.

이번 분석의 핵심 질문은 하나입니다. “현재 순위만 보고 가을야구 진출 팀을 판단해도 될까?”입니다.

## 데이터 설명

사용한 데이터는 팀 순위, 팀 타자 기록, 팀 투수 기록, 수비, 주루, 선수 기록, 선수 이동 현황입니다. 이 데이터를 그대로 쓰지 않고, 팀명 표준화, 숫자형 변환, 이닝 변환, 파생 지표 생성을 거쳐 분석용 마스터 테이블을 만들었습니다.

특히 팀 단위로 순위, 타격, 투수, 수비, 주루 데이터를 합친 `team_master`를 만들었고, 이 테이블을 모델 학습의 핵심 데이터로 사용했습니다.

## 전처리 설명

전처리에서는 먼저 팀명 표기를 통일했습니다. 그리고 승률, ERA, OPS, WHIP 같은 지표를 숫자형으로 변환했습니다. 투수 이닝처럼 `23 1/3` 형태로 들어온 값은 계산 가능한 소수형으로 변환했습니다.

이후 OPS, BABIP, ISO, K/9, BB/9, FIP, 득실차, 기대승률 같은 파생 지표를 만들었습니다.

## 모델 설명

모델은 Logistic Regression, Random Forest, Gradient Boosting 세 가지를 사용했습니다. 각각의 모델이 예측한 확률을 평균내는 앙상블 방식을 적용했습니다.

학습 데이터는 2022년부터 2025년까지의 완료 시즌이고, 예측 대상은 2026년 현재 시즌입니다. 목표 변수는 최종 순위 5위 이내, 즉 가을야구 진출 여부입니다.

## 결과 설명

최종 예측 결과, 가을야구 진출 가능성이 높게 나온 팀은 KT, LG, SSG, 삼성입니다.

KT는 {fmt_pct(pred.iloc[0]["가을야구_확률(%)"])}, LG는 {fmt_pct(pred.iloc[1]["가을야구_확률(%)"])}, SSG는 {fmt_pct(pred.iloc[2]["가을야구_확률(%)"])}, 삼성은 {fmt_pct(pred.iloc[3]["가을야구_확률(%)"])}로 높게 나왔습니다.

반면 현재 5위인 KIA는 예측 확률이 {fmt_pct(kia["가을야구_확률(%)"])}로 낮게 나왔습니다. 이 부분이 가장 중요한 인사이트입니다. 현재 순위는 5위지만, ERA와 WHIP 등 세부 지표가 약해서 모델은 위험 신호로 판단했습니다.

한화는 현재 6위지만 {fmt_pct(hanwha["가을야구_확률(%)"])}로 미진출 팀 중 가장 높았습니다. 그래서 완전히 탈락권이라기보다 경계선 팀이라고 볼 수 있습니다.

## 인사이트 설명

이번 분석에서 중요한 것은 현재 순위보다 승률의 질입니다. 홈에서만 강한 팀보다 원정에서도 성적을 유지하는 팀이 더 안정적으로 평가됐고, ERA와 WHIP 같은 투수 지표도 가을야구 가능성에 큰 영향을 줬습니다.

또 4월 순위만으로 최종 가을야구를 예측하는 정확도는 약 60% 수준이었습니다. 그래서 현재 순위만으로 판단하기보다, OPS, ERA, WHIP, 득실차, 기대승률을 함께 봐야 합니다.

## 결론

결론적으로 이 프로젝트는 단순 순위표가 아니라 세부 전력 지표를 기반으로 가을야구 가능성을 판단한 분석입니다.

현재 순위와 모델 예측은 완전히 같지 않았습니다. 특히 KIA처럼 현재 5위라도 세부 지표가 약하면 위험 신호가 나오고, 한화처럼 5위 밖이어도 추격 가능성이 남은 팀이 있습니다.

따라서 이 모델은 현재 순위표를 보완해서, 팀의 실제 전력과 향후 가능성을 판단하는 데 의미가 있습니다.
"""


def write_docx() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KBO 2026 가을야구 진출 예측 설명 자료")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("분석 기준: 2026년 4월 24일 진행중 시즌 데이터")
    doc.add_paragraph("사용 파일: KBO_가을야구_예측_전처리_v2.ipynb, 2026_가을야구_예측결과.csv, team_master_2022_2026.csv")

    for heading, body in [
        ("1. 분석 목적", "현재 순위만으로는 가을야구 진출 가능성을 설명하기 어렵기 때문에, 팀 타격·투수·수비·주루·홈/원정 성적을 통합해 2026년 가을야구 가능성을 예측했다."),
        ("2. 데이터 전처리", "팀명 표준화, 숫자형 변환, 투수 이닝 변환, OPS/BABIP/ISO/K9/BB9/FIP/득실차/기대승률 등 파생 지표 생성을 수행했다."),
        ("3. 모델 구조", "Logistic Regression, Random Forest, Gradient Boosting 세 모델의 예측 확률을 평균내는 앙상블 모델을 사용했다."),
        ("4. 핵심 인사이트", "현재 5위인 KIA는 세부 지표가 약해 낮은 확률이 나왔고, 6위 한화는 미진출 팀 중 가장 높은 확률로 경계선 팀으로 해석된다."),
    ]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)

    doc.add_heading("5. 2026 예측 결과", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["팀", "순위", "승패", "ERA", "OPS", "확률/예측"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for _, r in pred.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["팀명"])
        cells[1].text = str(int(r["순위"]))
        cells[2].text = f"{int(r['승'])}승 {int(r['패'])}패"
        cells[3].text = f"{float(r['ERA']):.2f}"
        cells[4].text = f"{float(r['OPS']):.3f}"
        cells[5].text = f"{float(r['가을야구_확률(%)']):.1f}% / {r['예측']}"

    doc.add_heading("6. 발표 결론", level=1)
    doc.add_paragraph(
        "가을야구 가능성은 현재 순위만으로 판단하기 어렵다. 승률의 안정성, 홈/원정 균형, 투수력, OPS, 득실차, 기대승률을 함께 봐야 한다. "
        "이 모델은 현재 순위표를 보완해 팀의 실제 전력과 향후 가능성을 설명하는 데 의미가 있다."
    )

    chart_paths = [
        ROOT / "02_rank_heatmap.png",
        ROOT / "06_feature_importance.png",
        ROOT / "07_playoff_prob.png",
        ROOT / "08_apr_vs_final.png",
    ]
    doc.add_heading("7. 참고 시각화", level=1)
    for path in chart_paths:
        if path.exists():
            doc.add_paragraph(path.name)
            doc.add_picture(str(path), width=Inches(5.8))

    doc.save(DOCX_PATH)


MD_PATH.write_text(md, encoding="utf-8")
SCRIPT_PATH.write_text(script, encoding="utf-8")
write_docx()

print(MD_PATH)
print(SCRIPT_PATH)
print(DOCX_PATH)
